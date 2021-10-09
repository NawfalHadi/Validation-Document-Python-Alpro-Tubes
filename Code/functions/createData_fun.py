from PyQt5.QtWidgets import *
import docx, os, shutil, datetime, string, time


def getDic(self, windowApp, textDic=QLabel):
    self.takeDic = QFileDialog(windowApp)
    opens = self.takeDic.getOpenFileName(self, 'Open File', 'D:/', 'Document (*.docx)')
    textDic.setText(opens[0])


def getNim(path):
    tempDoc = docx.Document(path)
    lenParDoc = len(tempDoc.paragraphs)
    tempNimList = []
    with open('Data/infoUser.txt') as user:
        for i in user:
            temp = i.split(',')
            tempNimList.append(temp[0])

    tempSorted = sorted(tempNimList)
    for i in tempSorted:
        for j in range(lenParDoc):
            if i in tempDoc.paragraphs[j].text:
                return i


def searchTitle(self, title, parag):
    thisTitle = title.split()
    print("Bismillah")
    if len(parag) > 15:
        print("bismillah..")
        cntr = 0
        print("bismillah..")
        for i in range(15):
            if len(parag[i].text) > 0:
                for j in thisTitle:
                    print("bismillah..")
                    if j in parag[i].text:
                        cntr += 1
                    if cntr > 4:
                        return [1, 'succes']
        if cntr > 0:
            return [1, 'succes']
        elif cntr < 0:
            return [0, 'Tolong gunakan judul dokumen yang ada di dalam file.']

    else:
        cntr = 0
        for i in parag:
            if len(i.text) > 0:
                for j in thisTitle:
                    if j in i.text:
                        cntr += 1
                    if cntr > 4:
                        return [1, 'succes']
        if cntr > 0:
            return [1, 'succes']
        elif cntr < 0:
            return [0, 'Tolong gunakan judul dokumen yang ada di dalam file.']


def searchNim(self, doc):
    docing = docx.Document(doc).paragraphs

    temp = []
    with open('Data/infoUser.txt', 'r')as nims:
        for nim in nims:
            temp.append(nim.split(',')[0])

    sorted(temp)

    count = 0
    for i in docing:
        for j in temp:
            for k in list(i.text):
                if k in string.digits:
                    count += 1
                    # Processing
                if len(j) == len(list(i.text)):
                    if j in i.text:
                        return [1, j]
    return [0, "Nim yang didalam dokumen tidak valid\n atau tidak dada didalam database"]


# This Night
def movingFile(self, title, nim, path):
    print(f"The title : {title}")
    print(f"yout nim : {nim}")

    vef_code = ''
    with open('Data/infoUser.txt') as user:
        for i in user:
            if nim in i:
                temp = i.split(',')
                vef_code = temp[2].strip()

    print(f"You Got Your pass : {vef_code}")
    if self.input_vef.text() == vef_code:
        print("bismillah")
        x = datetime.datetime.now()
        day = x.strftime("%y%m%d")

        newName = f"{path}/../{title}_{nim}_{day}.docx"
        print("Checking...")
        newFolder = f"Data/File/{nim}"
        createFolder(newFolder)
        print("Checking...")

        print(path)
        print(newName)

        os.rename(path, newName)
        print("Checking...")
        shutil.copy(newName, newFolder)
        print("Checking...")

        validating = 0
        documents = []
        with open(f"Data/User/{nim}.txt", 'r') as apper:
            for i in apper:
                validating += 1
                documents.append(i.split(',')[0])

        with open(f"Data/User/{nim}.txt", 'a') as appe:
            title2 = newName.split('/')[len(newName.split('/')) - 1]
            if validating != 0:
                if title2 in documents:
                    no = 2
                    while True:
                        new_title = f"{title2}({no})"
                        if new_title not in documents:
                            break
                        no += 1
                    appe.write(f"\n{new_title}, {day} ")
                    appe.close()
                else:
                    appe.write(f"\n{title2}, {day} ")
                    appe.close()
            elif validating == 0:
                appe.write(f"{title2}, {day} ")
                appe.close()

        print("Checking...")

        os.rename(newName, path)
        print("Checking...")
        print("alllhamdulillah")
        return [1, f'Dokumen {newName.split("/")[len(newName.split("/")) - 1]} sudah dimasukkan']
    elif self.input_vef.text() != vef_code:
        return [0, 'Verify code yang anda masukkan salah\n '
                   'tidak sesuai dengan nim yang ada di dokumen']  # Change to pop up error message
    else:
        print("Hah?!")


def createFolder(directory):
    try:
        if not os.path.exists(directory):
            os.makedirs(directory)
    except OSError:
        print('Error: Creating directory. ' + directory)
